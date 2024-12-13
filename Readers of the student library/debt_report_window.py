import sqlite3
from PyQt6.QtWidgets import QWidget, QVBoxLayout, QPushButton, QMessageBox, QTextEdit
from PyQt6.QtCore import QStandardPaths
from datetime import datetime
import os
from docx import Document

class DebtReportWindow(QWidget):
    def __init__(self, librarian_id):
        super().__init__()
        self.librarian_id = librarian_id
        self.setWindowTitle("Отчёт по долгам читателей")
        self.setFixedSize(700, 500)
        # Создание виджетов
        self.text_edit_report = QTextEdit(self)
        self.text_edit_report.setReadOnly(True)
        self.generate_report_button = QPushButton("Сгенерировать отчёт", self)
        self.generate_report_button.setStyleSheet("""
                    background-color: #FFA500;
                    border-radius: 5px;
                    color: white;
                    padding: 12px;
                    font-size: 18px;""")
        self.generate_report_button.clicked.connect(self.generate_report)
        # Кнопка для экспорта отчёта в Word
        self.export_to_word_button = QPushButton("Экспортировать в Word", self)
        self.export_to_word_button.setStyleSheet("""
                    background-color: #FFA500;
                    border-radius: 5px;
                    color: white;
                    padding: 12px;
                    font-size: 18px;""")
        self.export_to_word_button.clicked.connect(self.export_to_word)
        layout = QVBoxLayout(self)
        layout.addWidget(self.text_edit_report)
        layout.addWidget(self.generate_report_button)
        layout.addWidget(self.export_to_word_button)
        self.setLayout(layout)

    @staticmethod
    def log_operation(operation, id_librarian):
        desktop_path = QStandardPaths.writableLocation(QStandardPaths.StandardLocation.DesktopLocation)
        db_folder_path = os.path.join(desktop_path, "Library")
        db_path = os.path.join(db_folder_path, "library.db")

        try:
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()

            cursor.execute(
                "INSERT INTO Log (operation, id_librarian) VALUES (?, ?)",
                (operation, id_librarian)
            )

            conn.commit()

        except sqlite3.Error as e:
            print(f"Ошибка записи логов: {e}")

        finally:
            if 'conn' in locals():
                conn.close()

    def generate_report(self, librarian_id):
        self.log_operation("Генерация отчёта", self.librarian_id)
        try:
            desktop_path = QStandardPaths.writableLocation(QStandardPaths.StandardLocation.DesktopLocation)
            db_folder_path = os.path.join(desktop_path, "Library")
            db_path = os.path.join(db_folder_path, "library.db")
            with sqlite3.connect(db_path) as conn:
                cursor = conn.cursor()
                # Получаем текущую дату в формате строки
                current_date = datetime.now().strftime('%Y-%m-%d')

                # Читатели, которые вернули книги, но не оплатили штраф
                query_unpaid_fines = '''
                    SELECT r.surname, r.name, r.otchestvo, b.title, rec.fine 
                    FROM Records rec
                    JOIN Readers r ON rec.id_reader = r.id_reader
                    JOIN Books b ON rec.id_book = b.id_book
                    WHERE rec.date_return IS NOT NULL AND rec.fine > 0 AND rec.date_return <= ?;'''
                cursor.execute(query_unpaid_fines, (current_date,))
                unpaid_fines_data = cursor.fetchall()

                # Читатели, которые не вернули книги и их штрафы (добавлена дата выдачи)
                query_unreturned_books = '''
                    SELECT r.surname, r.name, r.otchestvo, b.title, rec.fine, rec.date_loan 
                    FROM Records rec
                    JOIN Readers r ON rec.id_reader = r.id_reader
                    JOIN Books b ON rec.id_book = b.id_book
                    WHERE rec.date_return IS NULL AND rec.fine > 0 AND rec.date_loan < ?;'''
                cursor.execute(query_unreturned_books, (current_date,))
                unreturned_books_data = cursor.fetchall()
            # Формирование отчёта
            report = ""
            report += "Штраф увеличивается каждый день для тех читателей, которые не возвращают книги!\n\n"  # Предупреждение
            if unpaid_fines_data:
                report += "Читатели, вернувшие книгу, но не оплатившие штраф:\n"
                for index, row in enumerate(unpaid_fines_data, start=1): # Нумерация с 1
                    fine_one = float(row[4])
                    report += f"{index}. {row[0]} {row[1]} {row[2]}  Книга: {row[3]}  Штраф: {fine_one:.2f} RUB.\n"
            else:
                report += "Нет читателей с неоплаченными штрафами.\n"
            report += "\n"

            if unreturned_books_data:
                report += "Читатели, не вернувшие книги и их текущий штраф:\n"
                for index, row in enumerate(unreturned_books_data, start=1):
                    fine = float(row[4])
                    report += f"{index}. {row[0]} {row[1]} {row[2]}  Книга: {row[3]}  Дата выдачи: {row[5]}  Штраф на текущий день: {fine:.2f} RUB.\n"
            else:
                report += "Нет читателей с не возвращёнными книгами.\n"
            # Выводим отчёт в текстовом поле
            self.text_edit_report.setText(report)
        except sqlite3.DatabaseError as db_error:
            # Обработка ошибок базы данных
            QMessageBox.critical(self, "Ошибка базы данных", f"Ошибка при работе с базой данных: {db_error}")
        except Exception as e:
            # Ловим все остальные ошибки
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка: {str(e)}")

    def export_to_word(self, librarian_id):
        # Получаем текст из QTextEdit (генерированный отчёт)
        report_text = self.text_edit_report.toPlainText()
        # Проверяем, если отчёт пустой
        if not report_text.strip():
            QMessageBox.warning(self, "Ошибка", "Нет данных для экспорта. Сначала сгенерируйте отчёт.")
            return
        # Создаём новый документ Word
        doc = Document()
        # Добавляем заголовок в документ
        doc.add_heading('Отчёт о долгах', 0)
        # Добавляем дату создания отчёта
        current_date = datetime.now().strftime('%Y-%m-%d')
        doc.add_paragraph(f'Дата создания отчёта: {current_date}\n')
        # Добавляем текст отчёта
        doc.add_paragraph(report_text)
        # Получаем путь к рабочему столу пользователя
        desktop_path = QStandardPaths.writableLocation(QStandardPaths.StandardLocation.DesktopLocation)
        folder_path = os.path.join(desktop_path, "Отчёты", "Долги")

        # Проверка, существует ли папка, если нет - создаём
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
            QMessageBox.information(self, "Информация", f"Папка '{folder_path}' была создана.")

        # Полный путь для сохранения файла
        save_path = os.path.join(folder_path, f"Отчёт_долги_{current_date}.docx")

        # Проверка, существует ли уже файл с таким именем
        if os.path.exists(save_path):
            # Если файл уже существует, показываем сообщение
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Информация")
            msg_box.setText(f"Файл '{os.path.basename(save_path)}' уже был экспортирован ранее в папку '{folder_path}'.")
            msg_box.setIcon(QMessageBox.Icon.Information)
            msg_box.setStandardButtons(QMessageBox.StandardButton.Ok)
            msg_box.exec()
            return
        try:
            doc.save(save_path)  # Сохраняем файл в указанное место
            # Получаем имя файла для отображения в сообщении
            file_name = os.path.basename(save_path)
            QMessageBox.information(self, "Успех", f"Отчёт успешно экспортирован в Word как '{file_name}' в папку '{folder_path}'")
            self.log_operation("Сохранение диаграммы в формате .docx", self.librarian_id)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при сохранении файла: {str(e)}")